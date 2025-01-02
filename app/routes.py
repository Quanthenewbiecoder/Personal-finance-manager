from flask import render_template, flash, redirect, url_for, request, jsonify, Response
from app import app, db
from app.models import User, Transaction, Category, DailyTotal
from flask_login import login_user, logout_user, current_user, login_required
from werkzeug.security import generate_password_hash, check_password_hash
import logging
from datetime import datetime, date as dt_date
from app.forms import EditTransactionForm, AddTransactionForm
import calendar
from collections import defaultdict
from sqlalchemy import func
from sqlalchemy.exc import IntegrityError
import csv
from docx import Document
import io


# Initialize logging
logging.basicConfig(level=logging.INFO)


@app.route('/')
@login_required
def index():
    # Fetch transactions for the current user
    transactions = Transaction.query.filter_by(user_id=current_user.id).order_by(Transaction.date.desc()).all()

    # Group transactions by year and month
    grouped_transactions = defaultdict(list)
    for transaction in transactions:
        month_year = transaction.date.strftime('%Y-%m')  # Format: 'YYYY-MM'
        grouped_transactions[month_year].append(transaction)

    return render_template('index.html', grouped_transactions=grouped_transactions)


@app.route('/download/<month_year>')
@login_required
def download_transactions(month_year):
    # Fetch transactions for the specified month and year
    transactions = Transaction.query.filter(
        Transaction.user_id == current_user.id,
        Transaction.date.like(f"{month_year}-%")
    ).all()

    # Check if transactions are available
    if not transactions:
        return Response(
            "No transactions found for the selected month.",
            mimetype='text/plain'
        )

    # Create a Word document
    doc = Document()
    doc.add_heading(f"Transactions for {month_year}", level=1)

    # Add a table to the document
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'

    # Add table headers
    header_cells = table.rows[0].cells
    header_cells[0].text = "ID"
    header_cells[1].text = "Description"
    header_cells[2].text = "Amount"
    header_cells[3].text = "Date"
    header_cells[4].text = "Category"

    # Add rows to the table
    for transaction in transactions:
        category = transaction.category.name if transaction.category else 'N/A'
        row_cells = table.add_row().cells
        row_cells[0].text = str(transaction.id)
        row_cells[1].text = transaction.description
        row_cells[2].text = f"{transaction.amount:.2f}"
        row_cells[3].text = transaction.date.strftime('%Y-%m-%d')
        row_cells[4].text = category

    # Save the document to an in-memory file
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    # Return the Word file as a response
    response = Response(file_stream.read(), mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response.headers.set('Content-Disposition', f'attachment; filename={month_year}_transactions.docx')
    return response


@app.route('/register', methods=['GET', 'POST'])
def register():
    if current_user.is_authenticated:
        return jsonify({"success": True, "message": "Already logged in.", "redirect_url": url_for('dashboard')}), 200

    if request.method == 'POST':
        username = request.form.get('username')
        email = request.form.get('email')
        password = request.form.get('password')
        confirm_password = request.form.get('confirm_password')

        # Ensure all fields are filled
        if not username or not email or not password or not confirm_password:
            return jsonify({"success": False, "message": "Please fill out all fields."}), 400

        # Check if passwords match
        if password != confirm_password:
            return jsonify({"success": False, "message": "Passwords do not match."}), 400

        # Check if password is the same as the username or email
        if password == username:
            return jsonify({"success": False, "message": "Password cannot be the same as the username."}), 400
        if password == email:
            return jsonify({"success": False, "message": "Password cannot be the same as the email."}), 400

        # Hash the password
        hashed_password = generate_password_hash(password)

        # Check if the hashed password already exists in the users table
        existing_user = User.query.filter(
            User.password_hash == hashed_password).first()
        if existing_user:
            return jsonify({"success": False, "message": "Password already in use. Please choose a different one."}), 400

        # Check if username or email already exists
        existing_user = User.query.filter(
            (User.email == email) | (User.username == username)).first()
        if existing_user:
            if existing_user.email == email:
                return jsonify({"success": False, "message": "Email already exists."}), 400
            if existing_user.username == username:
                return jsonify({"success": False, "message": "Username already exists."}), 400

        try:
            new_user = User(username=username, email=email,
                            password_hash=hashed_password)
            db.session.add(new_user)
            db.session.commit()
        except IntegrityError as e:
            db.session.rollback()
            logging.error(f"IntegrityError: {str(e)}")
            return jsonify({"success": False, "message": "An error occurred while registering the user."}), 500

        logging.info(f"New user registered: {username}")
        return jsonify({"success": True, "message": "Registration successful! Please log in.", "redirect_url": url_for('login')}), 200

    return render_template('register.html')


@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        # Redirect to overview if already logged in
        return redirect(url_for('overview'))

    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')

        if not username or not password:
            return jsonify({"success": False, "message": "Please enter both username and password."}), 400

        user = User.query.filter_by(username=username).first()

        if user and check_password_hash(user.password_hash, password):
            login_user(user)
            return jsonify({"success": True, "message": "Logged in successfully.", "redirect_url": url_for('overview')})
        else:
            return jsonify({"success": False, "message": "Invalid username or password."}), 400

    return render_template('login.html')


@app.route('/logout')
@login_required
def logout():
    username = current_user.username if current_user.is_authenticated else 'Anonymous'
    logout_user()
    flash('You have been logged out.', 'info')
    logging.info(f"User logged out: {username}")
    return redirect(url_for('index'))


@app.route('/add_transaction', methods=['GET', 'POST'])
@login_required
def add_transaction():
    form = AddTransactionForm()

    if form.validate_on_submit():
        description = form.description.data
        amount = form.amount.data
        date = form.date.data
        category_id = form.category.data

        if not description or not amount or not category_id:
            flash('Please fill out all fields.', 'warning')
            return redirect(url_for('add_transaction'))

        try:
            amount = float(amount)
            category_id = int(category_id)
        except ValueError:
            flash('Invalid amount or category ID.', 'danger')
            return redirect(url_for('add_transaction'))

        # Add the transaction to the database
        new_transaction = Transaction(description=description, amount=amount,
                                      date=date, category_id=category_id, user_id=current_user.id)
        db.session.add(new_transaction)

        # Update daily totals
        update_daily_totals(current_user.id, date)

        db.session.commit()

        flash('Transaction added successfully!', 'success')
        return redirect(url_for('dashboard'))

    categories = Category.query.all()
    return render_template('add_transaction.html', form=form, categories=categories)


def update_daily_totals(user_id, date):
    logging.info(f"Updating daily totals for user {user_id} on {date}")

    # Normalize date to YYYY-MM-DD format if it's not already a date object
    if isinstance(date, str):
        date = datetime.strptime(date, '%Y-%m-%d').date()
    elif isinstance(date, datetime):
        date = date.date()
    elif not isinstance(date, dt_date):
        raise ValueError(f"Unsupported date format: {type(date)}")

    # Retrieve transactions for the specified user and date
    transactions = Transaction.query.filter_by(user_id=user_id).filter(
        db.func.date(Transaction.date) == date).all()

    # Calculate totals
    total_income = sum(
        t.amount for t in transactions if t.category.name.lower() == 'income')
    total_expense = sum(
        t.amount for t in transactions if t.category.name.lower() == 'expense')
    total_savings = sum(
        t.amount for t in transactions if t.category.name.lower() == 'savings')
    net_total = total_income - total_expense + total_savings

    # Check if a daily total already exists for this date
    daily_total = DailyTotal.query.filter_by(
        user_id=user_id, date=date).first()

    if daily_total:
        # Update existing record
        daily_total.total_income = total_income
        daily_total.total_expense = total_expense
        daily_total.total_savings = total_savings
        daily_total.net_total = net_total
    else:
        # Create a new record
        new_daily_total = DailyTotal(
            user_id=user_id,
            date=date,
            total_income=total_income,
            total_expense=total_expense,
            total_savings=total_savings,
            net_total=net_total
        )
        db.session.add(new_daily_total)

    db.session.commit()
    logging.info("Daily totals updated and committed.")


@app.route('/dashboard', methods=['GET'])
@login_required
def dashboard():
    # Default to the current month and year if none provided
    today = datetime.today()
    selected_month = request.args.get('month', today.strftime('%m'))
    selected_year = request.args.get('year', today.strftime('%Y'))

    # Parse the selected month and year
    try:
        selected_year = int(selected_year)
        selected_month = int(selected_month)
    except ValueError:
        selected_month = today.month
        selected_year = today.year

    # Get month and year names for display
    selected_month_name = calendar.month_name[selected_month]
    months = {f'{month:02d}': f'{calendar.month_name[month]}' for month in range(1, 13)}
    years = list(range(2023, 2026))

    # Fetch transactions for the selected month and year for the current user
    start_date = f'{selected_year}-{selected_month:02d}-01'
    end_date = f'{selected_year}-{selected_month + 1:02d}-01' if selected_month < 12 else f'{selected_year + 1}-01-01'

    transactions = Transaction.query.filter(
        Transaction.user_id == current_user.id,  # Ensure transactions are for the current user
        Transaction.date >= start_date,
        Transaction.date < end_date
    ).order_by(Transaction.date.desc()).all()  # Order by newest first

    # Process transactions
    date_data = {}
    daily_totals = {}
    total_income = total_expense = total_savings = 0

    for trans in transactions:
        trans_date = trans.date.date()
        date_key = trans_date.strftime('%Y-%m-%d')

        if date_key not in date_data:
            date_data[date_key] = {'income': [], 'expense': [], 'savings': []}
            daily_totals[date_key] = {'income': 0, 'expense': 0, 'savings': 0}

        category_name = Category.query.get(trans.category_id).name.lower()  # Ensure lowercase comparison
        if category_name not in date_data[date_key]:
            date_data[date_key][category_name] = []

        date_data[date_key][category_name].append({
            'id': trans.id,
            'description': trans.description,
            'amount': trans.amount,
            'date': trans.date,
            'category': category_name
        })

        if category_name == 'income':
            daily_totals[date_key]['income'] += trans.amount
            total_income += trans.amount
        elif category_name == 'expense':
            daily_totals[date_key]['expense'] += trans.amount
            total_expense += trans.amount
        elif category_name == 'savings':
            daily_totals[date_key]['savings'] += trans.amount
            total_savings += trans.amount

    total_money = total_income - total_expense

    return render_template(
        'dashboard.html',
        total_money=total_money,
        total_savings=total_savings,
        date_data=date_data,
        daily_totals=daily_totals,
        selected_month=selected_month,
        selected_year=selected_year,
        selected_month_name=selected_month_name,
        months=months,
        years=years
    )

@app.route('/edit_transaction/<int:transaction_id>', methods=['GET', 'POST'])
@login_required
def edit_transaction(transaction_id):
    transaction = Transaction.query.filter_by(id=transaction_id, user_id=current_user.id).first_or_404()
    form = EditTransactionForm(
        obj=transaction,
        # Pass the current category type to the form
        category_type=transaction.category.name
    )

    if form.validate_on_submit():
        transaction.description = form.description.data
        transaction.amount = form.amount.data
        transaction.date = form.date.data
        transaction.category_id = form.category.data
        db.session.commit()
        flash('Transaction has been updated.', 'success')

        # Update daily totals after editing a transaction
        update_daily_totals(current_user.id, transaction.date)

        return redirect(url_for('dashboard'))

    return render_template('edit_transaction.html', form=form, transaction=transaction)


@app.route('/delete_transaction/<int:transaction_id>', methods=['POST'])
@login_required
def delete_transaction(transaction_id):
    transaction = Transaction.query.filter_by(id=transaction_id, user_id=current_user.id).first_or_404()

    # Delete the transaction
    db.session.delete(transaction)
    db.session.commit()

    flash('Transaction has been deleted.', 'success')

    # Update daily totals after deleting a transaction
    update_daily_totals(current_user.id, transaction.date)

    return redirect(url_for('dashboard'))


@app.route('/overview')
@login_required
def overview():
    # Define category IDs
    CATEGORY_ID_INCOME = 1
    CATEGORY_ID_EXPENSE = 2
    CATEGORY_ID_SAVINGS = 3

    # Calculate total income for all months for the current user
    total_income = db.session.query(func.sum(Transaction.amount)).filter(
        Transaction.user_id == current_user.id,
        Transaction.category_id == CATEGORY_ID_INCOME,
        Transaction.amount > 0
    ).scalar() or 0

    # Calculate total expenses for all months for the current user
    total_expenses = db.session.query(func.sum(Transaction.amount)).filter(
        Transaction.user_id == current_user.id,
        Transaction.category_id == CATEGORY_ID_EXPENSE
    ).scalar() or 0

    # Calculate total savings for all months for the current user
    total_savings = db.session.query(func.sum(Transaction.amount)).filter(
        Transaction.user_id == current_user.id,
        Transaction.category_id == CATEGORY_ID_SAVINGS,
        Transaction.amount > 0
    ).scalar() or 0

    # Calculate total money (income - expenses)
    total_money = total_income - total_expenses

    return render_template('overview.html', total_money=total_money, total_savings=total_savings)
